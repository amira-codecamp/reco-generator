from django.shortcuts import render
from django.http import HttpResponse
from wkhtmltopdf import render_pdf_from_template
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def index(request):
    context = {}
    return render(request, 'reco/templates/index.html', context)

def pdf(request):
    context = {}
    data = request.POST
    context['user'] = data.get('user')
    context['company'] = data.get('company')
    context['email'] = data.get('email')
    context['candidate'] = data.get('candidate')
    context['position'] = data.get('position')
    context['motivation'] = data.get('motivation')
    context['solving'] = data.get('solving')
    context['communication'] = data.get('communication')
    context['teamwork'] = data.get('teamwork')
    context['adaptability'] = data.get('adaptability')
    context['creativity'] = data.get('creativity')
    context['text'] = data.get('text')
    template_name = 'reco/templates/pdf.html'
    cmd_options={'orientation':'Portrait', 'page-size':'A4', 'margin-top':'80px', 'margin-bottom':'80px', 'margin-left':'80px', 'margin-right':'80px'}
    pdf_ = render_pdf_from_template(template_name, header_template=None, footer_template=None, context=context, cmd_options=cmd_options, request=request)
    response = HttpResponse(pdf_, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="reference-letter.pdf"'
    return response

def word(request):
    data = request.POST
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Helvetica'
    font.size = Pt(12)
    p = document.add_paragraph()
    p.style = document.styles['Normal']
    p.alignment = 3
    runner = p.add_run(data.get('user'))
    runner.bold = True
    font = runner.font
    font.color.rgb = RGBColor(88, 88, 88)
    p = document.add_paragraph()
    p.style = document.styles['Normal']
    p.alignment = 3
    runner = p.add_run(data.get('company'))
    runner.bold = True
    font = runner.font
    font.color.rgb = RGBColor(88, 88, 88)
    p = document.add_paragraph()
    p.style = document.styles['Normal']
    p.alignment = 3
    runner = p.add_run(data.get('email'))
    runner.bold = True
    font = runner.font
    font.color.rgb = RGBColor(88, 88, 88)
    p = document.add_paragraph()
    runner = p.add_run()
    runner.add_break()
    runner.add_break()
    p = document.add_paragraph()
    p.style = document.styles['Normal']
    p.alignment = 3
    runner = p.add_run('Subject: Recommendation for ' + data.get('candidate') + ' - ' + data.get('position'))
    runner.bold = True
    font = runner.font
    font.color.rgb = RGBColor(88, 88, 88)
    runner = p.add_run()
    runner.add_break()
    p = document.add_paragraph()
    p.style = document.styles['Normal']
    p.alignment = 3
    runner = p.add_run(data.get('text'))
    font = runner.font
    font.color.rgb = RGBColor(0, 0, 0)
    p = document.add_paragraph()
    runner = p.add_run()
    runner.add_break()
    matrix = document.add_table(rows=1, cols=2)
    row = matrix.rows[0].cells
    row[0].paragraphs[0].add_run('Skill').bold = True
    row[1].paragraphs[0].add_run('Level').bold = True
    level = ['Bad', 'Acceptable', 'Medium', 'Good', 'Very good']
    rate = ['1', '2', '3', '4', '5']
    color = [RGBColor(255, 133, 102), RGBColor(255, 173, 51), RGBColor(255, 255, 128), RGBColor(132, 225, 132), RGBColor(71, 209, 71)]
    row = matrix.add_row().cells
    row[0].text = 'Motivation'
    row[1].text = level[rate.index(data.get('motivation'))]
    row = matrix.add_row().cells
    row[0].text = 'Problem-solving'
    row[1].text = level[rate.index(data.get('solving'))]
    row = matrix.add_row().cells
    row[0].text = 'Communication'
    row[1].text = level[rate.index(data.get('communication'))]
    row = matrix.add_row().cells
    row[0].text = 'Teamwork'
    row[1].text = level[rate.index(data.get('teamwork'))]
    row = matrix.add_row().cells
    row[0].text = 'Adaptability'
    row[1].text = level[rate.index(data.get('adaptability'))]
    row = matrix.add_row().cells
    row[0].text = 'Creative thinking'
    row[1].text = level[rate.index(data.get('creativity'))]
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color[rate.index(data.get('motivation'))]))
    matrix.rows[1].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_2 = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color[rate.index(data.get('solving'))]))
    matrix.rows[2].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)
    shading_elm_3 = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color[rate.index(data.get('communication'))]))
    matrix.rows[3].cells[1]._tc.get_or_add_tcPr().append(shading_elm_3)
    shading_elm_4 = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color[rate.index(data.get('teamwork'))]))
    matrix.rows[4].cells[1]._tc.get_or_add_tcPr().append(shading_elm_4)
    shading_elm_5 = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color[rate.index(data.get('adaptability'))]))
    matrix.rows[5].cells[1]._tc.get_or_add_tcPr().append(shading_elm_5)
    shading_elm_6 = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color[rate.index(data.get('creativity'))]))
    matrix.rows[6].cells[1]._tc.get_or_add_tcPr().append(shading_elm_6)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=reference-letter.docx'
    document.save(response)
    return response
